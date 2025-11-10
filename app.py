# app.py
import streamlit as st
import pandas as pd
import google.generativeai as genai
import docx
import zipfile
import os
import io
import json
import re
import shutil
import tempfile
import time
from datetime import datetime
from google.api_core import exceptions
from PIL import Image, UnidentifiedImageError
import fitz  # PyMuPDF
import py7zr
import rarfile

# ==============================
# Streamlit Page Configuration
# ==============================
st.set_page_config(layout="wide", page_title="AI Grading Assistant")

# ==============================
# Helpers: File & Text Extraction
# ==============================

def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract visible text from a DOCX byte stream."""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"[DOCX read error] {e}"

def extract_images_from_docx_bytes(docx_bytes: bytes):
    """Extract images from a DOCX byte stream as PIL Images."""
    images = []
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data))
                    images.append(img)
                except UnidentifiedImageError:
                    continue
                except Exception:
                    continue
    except Exception:
        pass
    return images

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from a PDF byte stream using PyMuPDF."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        chunks = []
        for page in doc:
            chunks.append(page.get_text())
        return "\n".join(chunks)
    except Exception as e:
        return f"[PDF read error] {e}"

def read_code_from_ipynb(path: str) -> str:
    """Read concatenated code cells from a .ipynb file."""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            nb = json.load(f)
        cells = []
        for cell in nb.get("cells", []):
            if cell.get("cell_type") == "code":
                src = cell.get("source", [])
                if isinstance(src, list):
                    cells.append("".join(src))
                else:
                    cells.append(str(src))
        return "\n\n".join(cells)
    except Exception as e:
        return f"[IPYNB read error: {os.path.basename(path)}] {e}"

def read_text_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[TEXT read error: {os.path.basename(path)}] {e}"

def read_code_or_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in [".py", ".txt", ".md", ".json", ".csv", ".yaml", ".yml", ".ini", ".toml"]:
        return read_text_file(path)
    if ext == ".ipynb":
        return read_code_from_ipynb(path)
    return ""

def load_image_paths_and_docx_images(root_dir: str):
    """Collect image file paths in folder + extract images embedded in DOCX files."""
    pil_images = []

    # Images from folder structure
    IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"}
    for dirpath, _, filenames in os.walk(root_dir):
        for fn in filenames:
            ext = os.path.splitext(fn)[1].lower()
            if ext in IMAGE_EXTS:
                try:
                    img = Image.open(os.path.join(dirpath, fn))
                    pil_images.append(img)
                except Exception:
                    continue

    # Images embedded in DOCX files
    for dirpath, _, filenames in os.walk(root_dir):
        for fn in filenames:
            if fn.lower().endswith(".docx"):
                try:
                    with open(os.path.join(dirpath, fn), "rb") as f:
                        docx_bytes = f.read()
                    pil_images.extend(extract_images_from_docx_bytes(docx_bytes))
                except Exception:
                    continue

    return pil_images

# ==============================
# Archive Extraction (Nested)
# ==============================

def extract_archive(archive_path: str, dest: str):
    """Extract .zip, .7z, .rar archives into dest. Remove archive after extraction when possible."""
    ap = archive_path.lower()
    try:
        if ap.endswith(".zip"):
            with zipfile.ZipFile(archive_path, "r") as zf:
                zf.extractall(dest)
        elif ap.endswith(".7z"):
            with py7zr.SevenZipFile(archive_path, mode="r") as z:
                z.extractall(path=dest)
        elif ap.endswith(".rar"):
            # rarfile requires unrar/bsdtar on system; handle gracefully
            try:
                with rarfile.RarFile(archive_path) as rf:
                    rf.extractall(dest)
            except rarfile.RarCannotExec as e:
                st.warning(f"Unable to extract RAR (needs 'unrar' or 'bsdtar' installed): {os.path.basename(archive_path)}. {e}")
            except Exception as e:
                st.warning(f"RAR extraction failed for {os.path.basename(archive_path)}: {e}")
        else:
            return False
        try:
            os.remove(archive_path)
        except Exception:
            pass
        return True
    except Exception as e:
        st.warning(f"Archive extraction failed for {os.path.basename(archive_path)}: {e}")
        return False

def recursively_extract_archives(root_dir: str):
    """Find and extract all nested archives until none remain."""
    while True:
        archives = []
        for dirpath, _, filenames in os.walk(root_dir):
            for fn in filenames:
                if fn.lower().endswith((".zip", ".7z", ".rar")):
                    archives.append(os.path.join(dirpath, fn))
        if not archives:
            break
        for apath in archives:
            extract_archive(apath, os.path.dirname(apath))

# ==============================
# Directory Discovery
# ==============================

def discover_student_roots(extracted_root: str):
    """
    Heuristic:
    - If top-level contains multiple directories -> each is a student root.
    - If top-level contains files only -> treat the top-level as a single student's submission.
    - If there is a single directory -> dive in and repeat (to avoid unnecessary wrapper folders).
    """
    cur = extracted_root
    # Unwrap single-folder wrappers
    while True:
        items = [os.path.join(cur, x) for x in os.listdir(cur)]
        dirs = [x for x in items if os.path.isdir(x)]
        files = [x for x in items if os.path.isfile(x)]
        if len(dirs) == 1 and len(files) == 0:
            cur = dirs[0]
            continue
        break

    items = [os.path.join(cur, x) for x in os.listdir(cur)]
    dirs = [x for x in items if os.path.isdir(x)]
    files = [x for x in items if os.path.isfile(x)]

    if len(dirs) >= 1:
        return dirs
    else:
        return [cur]

# ==============================
# Gemini Setup & JSON Repair
# ==============================

@st.cache_resource
def get_gemini_model(_api_key: str):
    try:
        genai.configure(api_key=_api_key)
        available = [m.name for m in genai.list_models() if "generateContent" in m.supported_generation_methods]
        # Priorities — adjust as needed
        priorities = ["gemini-1.5-pro", "gemini-pro", "gemini-pro-vision"]
        pick = None
        for p in priorities:
            for m in available:
                if p in m:
                    pick = m
                    break
            if pick:
                break
        if not pick:
            st.error("No suitable Gemini model found for your API key.")
            return None, None
        return genai.GenerativeModel(pick), pick
    except Exception as e:
        st.error(f"Gemini configuration error: {e}")
        return None, None

def strip_code_fences(s: str) -> str:
    return s.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()

def try_extract_json_block(s: str) -> str:
    """Extract first top-level JSON object using a simple bracket counter."""
    start = s.find("{")
    if start == -1:
        return s
    depth = 0
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start:i+1]
    return s  # Fallback

def sanitize_json_text(s: str) -> str:
    """Common fixes: remove code fences, extract object, remove trailing commas, fix smart quotes."""
    s = strip_code_fences(s)
    s = try_extract_json_block(s)

    # Replace fancy quotes with standard quotes
    s = s.replace("“", '"').replace("”", '"').replace("’", "'")

    # Remove trailing commas before closing braces/brackets
    s = re.sub(r",(\s*[}\]])", r"\1", s)

    # Ensure keys are quoted (best-effort)
    # Be conservative to avoid damaging content inside strings
    def quote_keys(match):
        key = match.group(1)
        if key.startswith('"') and key.endswith('"'):
            return match.group(0)
        return f'"{key}":'
    s = re.sub(r"(?m)^\s*([A-Za-z_][A-Za-z0-9 _-]*)\s*:", quote_keys, s)

    return s

def parse_gemini_json(text: str):
    """Parse Gemini JSON with resilient fallback."""
    try:
        return json.loads(text)
    except Exception:
        pass
    try:
        return json.loads(sanitize_json_text(text))
    except Exception:
        # Last resort: find first {...} and load
        try:
            blk = try_extract_json_block(text)
            return json.loads(blk)
        except Exception:
            return None

# ==============================
# Rubric (from your brief)
# ==============================

RUBRIC_TEXT = """
Grade
Grade:
Rubric
Functional Requirements
No genuine attempt to define functional requirements
0 points
Fewer than 8 FRs, vague or not clearly mapped to Q1–Q5.
0.5 points
8–15 FRs, some mapping to Q1–Q5
1 points
8–15 clear, testable FRs, well-mapped to Q1–Q5.
2 points

Loading and Processing Data
No genuine attempt to load or process data.
0 points
Loads data but fails to handle invalid data, cancelled buses, or date/time formats correctly.
1 points
Fully correct parsing, including “none”, date/time, and cancelled buses.
3 points

Correctness – Q1–Q5 Implementation
No genuine attempt to implement requirements.
0 points
1–2 questions implemented with major issues.
1 points
3–4 questions implemented correctly
2 points
All 5 questions implemented correctly and appropriately.
3 points

Correctness - Output Accuracy
Outputs missing or clearly incorrect.
0 points
Outputs present but incorrect or inconsistent
1 points
Mostly accurate outputs, minor errors.
2 points
Outputs consistently accurate and match expectations
3 points

Visualization
No visualization provided.
0 points
Visualization present but unclear, incorrect, or unhelpful.
1 points
Clear visualization, but lacks labels or readability
2 points
Clear, labelled, readable charts that help answer at least one question.
3 points

User Interface
No genuine attempt
0 points
Interface present but confusing or incomplete
1 points
Usable interface with minor issues.
2 points
Clean and usable interface.
3 points

Code Quality - Organization
No meaningful structure; code is disorganized or hard to follow.
0 points
Code is reasonably well organized
0.5 points
Code is modular, logically organized, and easy to navigate
1 points

Code Quality - Naming conventions
Variable, function, or class names are unclear, inconsistent, or misleading.
0 points
Some meaningful names used, but inconsistently or with poor clarity
0.5 points
Names are meaningful and consistent
1 points

Code Quality - comments
No comments, code is undocumented.
0 points
Minimal or unclear comments that don’t aid understanding.
0.5 points
Use of comments to explain logic and structure as necessary.
1 points

Test Plan and Strategy
No test plan provided.
0 points
Basic plan, lacks coverage or mix of test types.
1.5 points
Clear strategy that will be able to effectively test the requirements
3 points

Coverage - Q1–Q5 Evidence
No evidence of coverage.
0 points
Evidence for 1–2 questions.
1 points
Evidence for 3–4 questions.
4 points
Evidence that the tool can answer each question
5 points

Manual Test Cases – Completeness
No manual test cases.
0 points
Inappropriate choice of manual test cases
1 points
Mostly correct test cases, missing one or two required components
2 points
Complete set with test ID, FRs, steps, expected results.
3 points

Manual Test Cases – Evidence
No evidence
0 points
Screenshots or logs missing or unclear.
1 points
Clear evidence (e.g. screenshots) for each test case.
2 points

Automated Tests
No automated tests
0 points
Tests described but not runnable or evidenced.
2 points
Runnable tests with partial evidence.
4 points
Clear descriptions, runnable, with evidence of execution.
5 points

Reproducibility
No instructions to run or reproduce tests.
0 points
Mostly clear steps to run tests.
1 points
Clear steps to run tests.
2 points

Deduction - Statement of Completion
Not included
-1 points
Included!
0 points

Deduction - Usage of AI statement
Not included
-1 points
Included. Statement says AI is used, but no detail about how
-0.5 points
Included - no AI was used
0 points
Included - AI used, and description provided of how, including prompts as necessary
0 points

Deduction - Statement of Assistance
Not included
-1 points
Included!
0 points
"""

# Structured rubric metadata to remind Gemini about max scores
RUBRIC_MAX = {
    "Functional Requirements": 2,
    "Loading and Processing Data": 3,
    "Correctness – Q1–Q5 Implementation": 3,
    "Correctness - Output Accuracy": 3,
    "Visualization": 3,
    "User Interface": 3,
    "Code Quality - Organization": 1,
    "Code Quality - Naming conventions": 1,
    "Code Quality - comments": 1,
    "Test Plan and Strategy": 3,
    "Coverage - Q1–Q5 Evidence": 5,
    "Manual Test Cases – Completeness": 3,
    "Manual Test Cases – Evidence": 2,
    "Automated Tests": 5,
    "Reproducibility": 2,
    # Deductions are negative maxima
    "Deduction - Statement of Completion": -1,
    "Deduction - Usage of AI statement": -1,  # note: can be -0.5/0/ -1 based on rubric
    "Deduction - Statement of Assistance": -1,
}

RUBRIC_ORDER = list(RUBRIC_MAX.keys())

def rubric_table_df():
    rows = []
    for k, v in RUBRIC_MAX.items():
        rows.append({"Criterion": k, "Max Score": v})
    return pd.DataFrame(rows)

# ==============================
# Gemini Prompting
# ==============================

def build_prompt(context: str, strictness_instruction: str, student_folder_name: str, report_text: str, code_blobs: str, images_present: bool) -> list:
    rubric_json_hint = json.dumps(RUBRIC_MAX, indent=2)
    today_str = datetime.now().strftime("%Y-%m-%d")

    prompt = [
        "You are an expert university programming tutor and examiner.",
        strictness_instruction,
        (
            "Grade the student submission *strictly according to the rubric* provided below. "
            "The assignment may include code (Python, notebooks), a report (DOCX/PDF), images, and test evidence. "
            "Half of the marks emphasize testing and reproducibility. "
            "If content is missing, score zero for the relevant criteria."
        ),
        f"Date: {today_str}",
        "\nASSIGNMENT CONTEXT:\n" + (context or "No extra context provided."),
        "\nRUBRIC (full text):\n" + RUBRIC_TEXT,
        "\nRUBRIC MAX SCORES (machine-readable):\n" + rubric_json_hint,
        f"\nSTUDENT SUBMISSION FOLDER: {student_folder_name}",
        "\nREPORT TEXT:\n" + (report_text if report_text else "[No report text found]"),
        "\nCODE (combined):\n```python\n" + (code_blobs if code_blobs else "# No code detected\n") + "\n```",
        f"\nIMAGES PRESENT: {'Yes' if images_present else 'No'}",
        "\nREQUIREMENTS:",
        "- For each rubric criterion, assign a 'score' that is one of the allowed marks implied by the description, capped by the Max Score. Do not invent new criteria.",
        "- Apply deductions exactly for the 'Deduction - ...' items when missing; use negative scores up to their minimum.",
        "- Justify briefly (1–2 sentences) per criterion, citing concrete evidence from the submission.",
        "- Then provide a single concise overall feedback paragraph (2–4 sentences).",
        "- Return ONLY a single JSON object with this exact schema:",
        """
{
  "grading_summary": [
    { "criterion": "Functional Requirements", "score": 0.0, "max_score": 2, "justification": "..." },
    ...
  ],
  "overall_feedback": "..."
}
""",
        "- IMPORTANT: Ensure every criterion from the following list appears exactly once in 'grading_summary' and in this order:\n" + json.dumps(RUBRIC_ORDER),
        "- Ensure 'score' is numeric; 'max_score' must match the 'RUBRIC MAX SCORES' above for the criterion.",
        "- Do not include any text before/after the JSON object."
    ]
    return prompt

def grade_with_gemini(model, prompt_parts):
    if not model:
        return None, "Gemini model not available."
    # Retry with exponential backoff for quota
    delay = 5
    for attempt in range(3):
        try:
            resp = model.generate_content(prompt_parts, request_options={"timeout": 180})
            txt = resp.text or ""
            parsed = parse_gemini_json(txt)
            if not parsed:
                return None, "Failed to parse JSON from model output."
            return parsed, None
        except exceptions.ResourceExhausted as e:
            if attempt < 2:
                time.sleep(delay)
                delay *= 2
                continue
            return None, f"Rate limit / quota error: {e}"
        except Exception as e:
            return None, f"Unexpected error: {e}"
    return None, "Unknown error."

# ==============================
# Scoring Utilities
# ==============================

def flatten_result(student_name: str, gemini_json: dict):
    """Flatten grading JSON into a single dict row."""
    row = {"Student Folder": student_name}
    total = 0.0
    max_total = 0.0
    for item in gemini_json.get("grading_summary", []):
        crit = item.get("criterion", "").strip()
        score = float(item.get("score", 0))
        max_score = float(item.get("max_score", 0))
        just = item.get("justification", "")
        row[f"{crit} Score"] = score
        row[f"{crit} Justification"] = just
        total += score
        max_total += max_score
    row["Total Score"] = total
    row["Max Score"] = max_total if max_total > 0 else sum(abs(v) for v in RUBRIC_MAX.values())
    row["Percent"] = round((row["Total Score"] / row["Max Score"]) * 100, 2) if row["Max Score"] else 0.0
    row["Grade Band"] = grade_band(row["Percent"])
    row["Overall Feedback"] = gemini_json.get("overall_feedback", "")
    return row

def grade_band(percent: float) -> str:
    # Typical AU bands, adjust if needed
    if percent >= 80:
        return "HD"
    if percent >= 70:
        return "D"
    if percent >= 60:
        return "C"
    if percent >= 50:
        return "P"
    return "F"

# ==============================
# Core: Process one student folder
# ==============================

def compile_submission_texts(student_dir: str):
    """
    Return (report_text, code_text, images_present) for the student folder.
    - report_text: concatenation of all readable report-like files (.docx/.pdf/.txt/.md)
    - code_text: concatenation of code/notebook/exportable text (.py/.ipynb/.json/.csv)
    - images_present: whether any images exist (from folder or embedded in docx)
    """
    report_chunks = []
    code_chunks = []

    # Read DOCX/PDF/Text for report
    for dirpath, _, filenames in os.walk(student_dir):
        for fn in filenames:
            p = os.path.join(dirpath, fn)
            ext = os.path.splitext(fn)[1].lower()
            try:
                if ext == ".docx":
                    with open(p, "rb") as f:
                        b = f.read()
                    report_chunks.append(extract_text_from_docx_bytes(b))
                elif ext == ".pdf":
                    with open(p, "rb") as f:
                        b = f.read()
                    report_chunks.append(extract_text_from_pdf_bytes(b))
                elif ext in [".txt", ".md"]:
                    report_chunks.append(read_text_file(p))
            except Exception as e:
                report_chunks.append(f"[Read error: {fn}] {e}")

    # Read code-like / structured text
    for dirpath, _, filenames in os.walk(student_dir):
        for fn in filenames:
            p = os.path.join(dirpath, fn)
            ext = os.path.splitext(fn)[1].lower()
            try:
                if ext in [".py", ".ipynb", ".json", ".csv", ".yaml", ".yml", ".ini", ".toml"]:
                    code_chunks.append(read_code_or_text(p))
            except Exception as e:
                code_chunks.append(f"[Read error: {fn}] {e}")

    # Gather images (folder + embedded in DOCX)
    pil_images = load_image_paths_and_docx_images(student_dir)
    images_present = len(pil_images) > 0

    report_text = "\n\n".join([c for c in report_chunks if c])
    code_text = "\n\n".join([c for c in code_chunks if c])

    # Bound extremely long payloads to keep token usage sane
    MAX_CHARS = 120_000
    if len(report_text) > MAX_CHARS:
        report_text = report_text[:MAX_CHARS] + "\n...[truncated]..."
    if len(code_text) > MAX_CHARS:
        code_text = code_text[:MAX_CHARS] + "\n...[truncated]..."

    return report_text, code_text, images_present

# ==============================
# Batch Grading from ZIP
# ==============================

def find_and_grade_assignments(zip_file_obj, rubric_text: str, context: str, model, strictness_instruction: str):
    all_rows = []
    with tempfile.TemporaryDirectory() as temp_dir:
        status = st.empty()

        # Extract top-level ZIP
        try:
            with zipfile.ZipFile(zip_file_obj, "r") as zf:
                zf.extractall(temp_dir)
        except Exception as e:
            st.error(f"Failed to extract main ZIP: {e}")
            return []

        # Recursively extract nested archives
        status.info("Extracting nested archives (.zip/.7z/.rar)...")
        recursively_extract_archives(temp_dir)

        # Discover student roots
        student_dirs = discover_student_roots(temp_dir)
        if not student_dirs:
            st.warning("No student folders found.")
            return []

        progress = st.progress(0.0)
        total = len(student_dirs)

        for i, sdir in enumerate(student_dirs):
            student_name = os.path.basename(sdir.rstrip(os.sep))
            status.info(f"Grading {i+1}/{total}: {student_name}")
            try:
                report_text, code_text, images_present = compile_submission_texts(sdir)
                # If literally nothing, skip
                if not report_text and not code_text:
                    st.warning(f"Skipping '{student_name}': no readable content detected.")
                else:
                    prompt_parts = build_prompt(
                        context=context,
                        strictness_instruction=strictness_instruction,
                        student_folder_name=student_name,
                        report_text=report_text,
                        code_blobs=code_text,
                        images_present=images_present,
                    )
                    result_json, err = grade_with_gemini(model, prompt_parts)
                    if err:
                        st.error(f"Gemini grading error for '{student_name}': {err}")
                    elif not result_json:
                        st.error(f"Gemini returned no parseable JSON for '{student_name}'.")
                    else:
                        row = flatten_result(student_name, result_json)
                        all_rows.append(row)
            except Exception as e:
                st.error(f"Processing failed for '{student_name}': {e}")

            progress.progress((i + 1) / total)

        status.success("Grading complete!")
    return all_rows

# ==============================
# UI
# ==============================

st.title("👨‍🏫 AI-Powered Assignment Grader (Rubric-Aligned)")

st.markdown("""
This tool grades a **batch of submissions** from a ZIP file.  
It handles **nested folders**, **nested archives** (.zip, .7z, .rar), parses **DOCX/PDF/IPYNB/PY**, and includes **images**.  
Scoring strictly follows your rubric, with strong emphasis on **testing & reproducibility**.
""")

with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Gemini API Key", type="password")

    st.header("Grading Style")
    strictness_level = st.selectbox("Strictness", ("Lenient", "Standard", "Strict", "Critical"))
    strictness_map = {
        "Lenient": "INSTRUCTION: Grade leniently. Reward effort and partial progress.",
        "Standard": "INSTRUCTION: Grade fairly and accurately per rubric.",
        "Strict": "INSTRUCTION: Grade strictly. Deduct for minor errors or weak evidence.",
        "Critical": "INSTRUCTION: Grade with a highly critical eye, as for a capstone/final-year project.",
    }
    strictness_instruction = strictness_map[strictness_level]

    st.header("Rubric (Used As-Is)")
    rubric_text = st.text_area("Rubric", value=RUBRIC_TEXT, height=300)

    st.header("Assignment Context (optional)")
    assignment_context = st.text_area("Context", value="", height=120)

    st.header("Upload")
    uploaded_zip = st.file_uploader("Upload a ZIP containing all student folders", type="zip")

st.subheader("Rubric Summary")
st.dataframe(rubric_table_df(), use_container_width=True)

st.header("📝 Grading Results")

if "results_df" not in st.session_state:
    st.session_state.results_df = None

col_go, col_clear = st.columns([1,1])
with col_go:
    run_btn = st.button("🚀 Grade Assignments", type="primary", use_container_width=True)
with col_clear:
    if st.button("🧹 Clear Results", use_container_width=True):
        st.session_state.results_df = None

if run_btn:
    if not api_key:
        st.error("Please enter your Gemini API Key.")
    elif not rubric_text.strip():
        st.error("Please paste the rubric.")
    elif not uploaded_zip:
        st.error("Please upload a ZIP file.")
    else:
        model, model_name = get_gemini_model(api_key)
        if model:
            with st.spinner(f"Grading with {model_name} ({strictness_level})..."):
                results = find_and_grade_assignments(
                    uploaded_zip,
                    rubric_text,
                    assignment_context,
                    model,
                    strictness_instruction
                )
                if results:
                    st.session_state.results_df = pd.DataFrame(results)
                    st.success("Grading completed successfully.")
                else:
                    st.session_state.results_df = None
                    st.warning("No results generated.")

if st.session_state.results_df is not None:
    df = st.session_state.results_df
    st.dataframe(df, use_container_width=True)

    @st.cache_data
    def to_csv_bytes(df_: pd.DataFrame):
        return df_.to_csv(index=False).encode("utf-8")

    @st.cache_data
    def to_excel_bytes(df_: pd.DataFrame):
        bio = io.BytesIO()
        with pd.ExcelWriter(bio, engine="openpyxl") as writer:
            df_.to_excel(writer, index=False, sheet_name="Grades")
        return bio.getvalue()

    csv_data = to_csv_bytes(df)
    xlsx_data = to_excel_bytes(df)

    c1, c2 = st.columns(2)
    with c1:
        st.download_button("📥 Download CSV", data=csv_data, file_name="grading_report.csv", mime="text/csv", use_container_width=True)
    with c2:
        st.download_button("📥 Download Excel", data=xlsx_data, file_name="grading_report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

else:
    st.info("Results will appear here once grading finishes.")

st.caption("""
Notes:
- For RAR extraction, your system may need `unrar` or `bsdtar`. If unavailable, RARs will be skipped with a warning.
- Consider using a single venv or global site-packages to avoid installing heavy libs (e.g., pandas) many times across student projects.
- If students used AI, ensure their submission includes a **Usage of AI statement** and **Statement of Assistance** to avoid deductions.
""")
